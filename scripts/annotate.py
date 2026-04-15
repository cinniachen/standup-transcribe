#!/usr/bin/env python3
"""
standup-transcribe 喜剧技巧标注脚本
基于文本行列表生成标注版 Word 文档

用法：
  python annotate.py segments.json output.docx --performer "鸟鸟" --title "主咖和TA的朋友们"

输入格式：
  segments.json = Whisper 转录结果，每项 {"start": 0.0, "end": 5.0, "text": "..."}
  或纯文本行列表 ["句子1", "句子2", ...]
"""

import sys, os, json, re
from collections import Counter

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
# 喜剧技巧分类定义（融合五维度分析框架）
# ============================================================
#
# 维度一：核心四要素（好笑/真实/原创/对话）—— AI 分析注释的评估基准，不作为标签
# 维度二：素材选择五特征（真实/个人化/宜人度/熟悉度/普适性/荒谬点）—— AI 分析注释的评估维度
# 维度三：结构逻辑（引子→论点→荒诞点→例子→分论点）—— 作为结构标签
# 维度四：观点提炼（观点=事实+负面情绪；困难/奇怪/害怕/愚蠢）—— 作为观点标签
# 维度五：加梗技术（呈现/比喻/类比/夸张/假设/混合/首尾呼应）—— 作为技巧标签

TAG_COLORS = {
    # ---- 维度三：结构逻辑标签 ----
    "引子":    RGBColor(0x26, 0xA6, 0x9A),   # 青绿 - 开场钩子，引入话题
    "论点":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 核心观点陈述
    "分论点":  RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 子论点，支持主论点
    "荒诞点":  RGBColor(0xEC, 0x40, 0x7A),   # 粉红 - 点出荒谬之处
    "例证":    RGBColor(0x8D, 0x6E, 0x63),   # 棕 - 用例子论证论点
    "铺垫":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 建立前提/背景（引子+论点的蓄力阶段）
    "收尾":    RGBColor(0x26, 0xA6, 0x9A),   # 青绿 - 收束/结束

    # ---- 维度四：观点标签 ----
    "观点":    RGBColor(0xFF, 0xB3, 0x00),   # 金 - 观点=事实+负面情绪
    "负面情绪": RGBColor(0x78, 0x90, 0x9C),  # 灰蓝 - 困难/奇怪/害怕/愚蠢
    "社会观察": RGBColor(0x78, 0x90, 0x9C),  # 灰蓝 - 社会议题转化为喜剧素材

    # ---- 维度五：加梗技术标签 ----
    "呈现":    RGBColor(0x66, 0xBB, 0x6A),   # 绿 - Act-out，演出来而非转述
    "比喻":    RGBColor(0x7E, 0x57, 0xC2),   # 紫 - "就像……"产生画面感
    "类比":    RGBColor(0x7E, 0x57, 0xC2),   # 紫 - 用另一事物解释荒谬感
    "夸张":    RGBColor(0xEC, 0x40, 0x7A),   # 粉红 - 局部夸大制造不合理
    "假设":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 动机假设/后果假设
    "混合":    RGBColor(0xAB, 0x47, 0xBC),   # 深紫 - 不相关事物联系产生错位
    "call-back": RGBColor(0xF5, 0xA6, 0x23), # 橙 - 首尾呼应，回收前面的梗
    "回旋镖":  RGBColor(0xF5, 0xA6, 0x23),   # 橙 - 把别人的话原封扔回去

    # ---- 通用技巧标签（跨维度） ----
    "笑点":    RGBColor(0xE5, 0x3E, 0x3E),   # 红 - 包袱释放点
    "punchline": RGBColor(0xE5, 0x3E, 0x3E), # 红 - 同笑点
    "反转":    RGBColor(0xAB, 0x47, 0xBC),   # 深紫 - 打破预期
    "二次反转": RGBColor(0xAB, 0x47, 0xBC),  # 深紫 - 反转之后再反转
    "自嘲":    RGBColor(0x66, 0xBB, 0x6A),   # 绿 - 自我调侃
    "递进":    RGBColor(0x26, 0xC6, 0xDA),   # 青 - 层层升级
    "排比":    RGBColor(0x26, 0xC6, 0xDA),   # 青 - 节奏递进
    "金句":    RGBColor(0xFF, 0xB3, 0x00),   # 金 - 点睛之笔
    "观察":    RGBColor(0x8D, 0x6E, 0x63),   # 棕 - 生活细节提炼
    "暗喻":    RGBColor(0x7E, 0x57, 0xC2),   # 紫 - 比喻的隐式用法
    "反讽":    RGBColor(0x45, 0x5A, 0x64),   # 深灰 - 正话反说
    "meta笑点": RGBColor(0xFF, 0xB3, 0x00),  # 金 - 元喜剧
    "解构":    RGBColor(0xAB, 0x47, 0xBC),   # 深紫 - 解构常见概念
    "误导":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 引导到错误方向
    "悬念":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 制造期待
    "对比":    RGBColor(0x78, 0x90, 0x9C),   # 灰蓝 - 对照制造落差
    "极简收束": RGBColor(0x26, 0xA6, 0x9A),  # 青绿 - 极短收尾，节奏反差
    "钩子":    RGBColor(0x26, 0xA6, 0x9A),   # 青绿 - 开场钩子
    "修正/笑点": RGBColor(0xE5, 0x3E, 0x3E),# 红 - 修正前面的铺垫后出笑点
    "暗贬":    RGBColor(0x42, 0xA5, 0xF5),   # 蓝 - 隐性贬低
    "角色翻转": RGBColor(0xAB, 0x47, 0xBC),  # 深紫
    "视角翻转": RGBColor(0xAB, 0x47, 0xBC),  # 深紫
    "性别反转": RGBColor(0xAB, 0x47, 0xBC),  # 深紫
    "对称":    RGBColor(0x26, 0xC6, 0xDA),   # 青
    "矛盾修辞": RGBColor(0xAB, 0x47, 0xBC),  # 深紫
    "谐音梗":  RGBColor(0xEC, 0x40, 0x7A),   # 粉红
    "文化焦虑": RGBColor(0x78, 0x90, 0x9C),  # 灰蓝
    "主题升华": RGBColor(0xFF, 0xB3, 0x00),  # 金
    "黑色幽默": RGBColor(0x45, 0x5A, 0x64),  # 深灰

    # ---- 复合标签（AI 手动标注时使用） ----
    "铺垫+暗喻": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+排比": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+误导": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+观察": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+悬念": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+暗贬": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+反转": RGBColor(0x42, 0xA5, 0xF5),
    "铺垫+对比": RGBColor(0x42, 0xA5, 0xF5),
    "论点+荒诞点": RGBColor(0xEC, 0x40, 0x7A),
    "例证+呈现": RGBColor(0x66, 0xBB, 0x6A),
    "假设+夸张": RGBColor(0xEC, 0x40, 0x7A),
    "观点+自嘲": RGBColor(0x66, 0xBB, 0x6A),
    "比喻+类比": RGBColor(0x7E, 0x57, 0xC2),
    "call-back+回旋镖": RGBColor(0xF5, 0xA6, 0x23),
    "反转+二次反转": RGBColor(0xAB, 0x47, 0xBC),
}


def get_tag_color(tag_str):
    """取标签颜色"""
    if tag_str in TAG_COLORS:
        return TAG_COLORS[tag_str]
    first = tag_str.split('+')[0].strip()
    if first in TAG_COLORS:
        return TAG_COLORS[first]
    return RGBColor(0x78, 0x90, 0x9C)


# ============================================================
# 技巧检测：基于关键词/模式的自动标注
# ============================================================

# 笑点关键词和模式
PUNCH_PATTERNS = [
    r'^(所以|但|其实|结果|没想到|谁知道|后来|直到)',
    r'(没人|没有人|没有谁)(问|理|管|关心)',
    r'(到了不用|因为心已经死了|谢谢大家)',
    r'对不起',
    r'但我也能理解',
    r'我们盼望',
    r'(笑点|punchline|反转|吐槽|调侃|嘲讽|讽刺)',
]

SETUP_PATTERNS = [
    r'^(其实|说起来|你们知道|很多人|有一个|上次|昨天|前两天)',
    r'(最近|前阵子|有一年|那个时候)',
]

SELF_DEPREC_PATTERNS = [
    r'(我|我自己|我的)(也|就|才|就是|只是)',
    r'(没人|没有人|没有谁)(给我|问我|找我|关心我)',
    r'(普通|一般|正常|平凡)',
    r'(孤独|一个人|单身|社恐|内向|尴尬)',
]

CALLBACK_PATTERNS = [
    r'(回到|刚才|前面|之前|我们|那个|还记得)',
    r'(黄圣依|张萌|张绍刚|张伟丽|饶雪漫|李诞)',  # 如果嘉宾名字重复出现
]


def detect_technique(text, all_lines, current_idx, performer_name=''):
    """
    自动检测一段文本使用的喜剧技巧。
    返回 (标签, 注释说明) 或 None
    """
    text = text.strip()
    if not text:
        return None

    # 检查前面3行是否提到了相同关键词（call-back检测）
    prev_keywords = set()
    for j in range(max(0, current_idx - 8), current_idx):
        prev_text = all_lines[j].strip() if isinstance(all_lines[j], str) else all_lines[j].get('text', '')
        # 提取人名
        names = re.findall(r'([\u4e00-\u9fff]{2,4})', prev_text)
        for n in names:
            if n in text and n != performer_name and len(n) >= 2:
                prev_keywords.add(n)

    tags = []
    notes = []

    # 1. 检测自嘲
    for p in SELF_DEPREC_PATTERNS:
        if re.search(p, text):
            tags.append("自嘲")
            notes.append("把自己作为笑料的一部分，降低攻击性、建立亲和力")
            break

    # 2. 检测笑点特征
    is_punch = False
    for p in PUNCH_PATTERNS:
        if re.search(p, text):
            is_punch = True
            break
    # 极短句子（<=8字）往往是 punchline
    if len(text) <= 8 and not re.match(r'^(嗯|啊|那个|就是|对|好的|大家好)', text):
        is_punch = True
        tags.append("极简收束")
        notes.append("极简收尾，前文长铺垫后的短句 punch，制造节奏反差")

    if is_punch and "极简收束" not in tags:
        # 检查是否是反转
        if any(k in text for k in ['但', '其实', '结果', '没想到', '谁知']):
            tags.append("反转")
            notes.append("打破观众预期，制造意外")
        else:
            tags.append("笑点")
            notes.append("包袱释放点")

    # 3. 检测铺垫
    if not is_punch:
        for p in SETUP_PATTERNS:
            if re.search(p, text):
                tags.append("铺垫")
                break
        # 叙述性长句通常是铺垫
        if len(text) > 20 and "铺垫" not in tags:
            tags.append("铺垫")
            notes.append("建立场景或提供背景信息，为后续笑点蓄力")

    # 4. Call-back 检测
    if prev_keywords:
        tags.append("call-back")
        names_str = "、".join(prev_keywords)
        notes.append(f"回扣前面提到的「{names_str}」，串联前文制造'原来如此'的满足感")

    # 5. 排比检测（连续3+行相似句式）
    if current_idx >= 2:
        prev1 = all_lines[current_idx - 1].strip() if isinstance(all_lines[current_idx - 1], str) else all_lines[current_idx - 1].get('text', '')
        prev2 = all_lines[current_idx - 2].strip() if isinstance(all_lines[current_idx - 2], str) else all_lines[current_idx - 2].get('text', '')
        # 简单检测：三句话长度相近
        if abs(len(text) - len(prev1)) < 5 and abs(len(prev1) - len(prev2)) < 5:
            tags.append("排比")
            notes.append("排比句式，节奏递进增强力度")

    # 6. 金句检测（包含哲理或点睛之笔）
    quotable = [
        r'疼了还能呻吟', r'心已经死了', r'情绪稳定', r'无病呻吟',
        r'时光机', r'长痛不如短痛', r'仁慈',
    ]
    for q in quotable:
        if q in text:
            tags.append("金句")
            notes.append("点睛之笔，简洁有力，记忆点极强")
            break

    # 7. 递进检测
    escalation_markers = ['不如', '更', '甚至', '连', '不仅', '而且']
    count = sum(1 for m in escalation_markers if m in text)
    if count >= 2:
        if "递进" not in tags:
            tags.append("递进")
            notes.append("多级递进，层层升级制造笑点叠加效果")

    # 8. 反讽/社会观察
    social_markers = ['中年', '女性', '舆论', '性别', '身材', '体重', '年龄', '社会']
    if any(m in text for m in social_markers):
        if "社会观察" not in tags:
            tags.append("社会观察")
            notes.append("嵌入社会议题，笑完之后有思考空间")

    if not tags:
        return None

    tag_str = " + ".join(tags)
    note_str = "；".join(notes) if notes else ""
    return (tag_str, note_str)


# ============================================================
# 生成标注版 Word
# ============================================================

def generate_annotated_word(text_lines, title='喜剧技巧标注版', output_path=None,
                             performer='', custom_annotations=None):
    """
    生成喜剧技巧标注版 Word 文档。

    参数：
      text_lines: 文本行列表（字符串列表 或 Whisper segments 字典列表）
      title: 文档标题
      output_path: 输出路径（默认自动生成）
      performer: 表演者名字（用于 call-back 检测）
      custom_annotations: 自定义标注列表 [(文本, 标签, 注释), ...]（可选，会覆盖自动检测）
    """
    if not text_lines:
        print("错误：没有文本内容")
        return None

    # 统一为纯文本列表
    lines = []
    for item in text_lines:
        if isinstance(item, str):
            lines.append(item.strip())
        elif isinstance(item, dict):
            lines.append(item.get('text', '').strip())
        else:
            lines.append(str(item).strip())
    lines = [l for l in lines if l]

    if not output_path:
        # 自动生成文件名
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)[:30]
        output_path = os.path.expanduser(f'~/Downloads/{safe_title}_标注.docx')

    # 如果有自定义标注，使用自定义标注；否则自动检测
    if custom_annotations:
        segments = custom_annotations
    else:
        segments = []
        for i, line in enumerate(lines):
            result = detect_technique(line, lines, i, performer)
            if result:
                tag, note = result
                segments.append((line, tag, note))

    # ---- 创建文档 ----
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'PingFang SC'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # 标题
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('逐段标注 · 技巧拆解 · 结构分析')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)

    # 图例（按五维度分组）
    doc.add_paragraph('')
    legend = doc.add_paragraph()
    lt = legend.add_run('【图例说明 · 五维度标注体系】')
    lt.font.bold = True
    lt.font.size = Pt(10)
    legend.add_run('\n')

    # 维度三：结构逻辑
    r = legend.add_run('  ▸ 结构逻辑：')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    struct_items = [
        ("引子", RGBColor(0x26, 0xA6, 0x9A)),
        ("论点", RGBColor(0x42, 0xA5, 0xF5)),
        ("荒诞点", RGBColor(0xEC, 0x40, 0x7A)),
        ("例证", RGBColor(0x8D, 0x6E, 0x63)),
        ("铺垫", RGBColor(0x42, 0xA5, 0xF5)),
        ("收尾", RGBColor(0x26, 0xA6, 0x9A)),
    ]
    for text, color in struct_items:
        r = legend.add_run(f' [{text}]')
        r.font.size = Pt(9)
        r.font.color.rgb = color
    legend.add_run('\n')

    # 维度四：观点
    r = legend.add_run('  ▸ 观点提炼：')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    view_items = [
        ("观点", RGBColor(0xFF, 0xB3, 0x00)),
        ("负面情绪", RGBColor(0x78, 0x90, 0x9C)),
        ("社会观察", RGBColor(0x78, 0x90, 0x9C)),
    ]
    for text, color in view_items:
        r = legend.add_run(f' [{text}]')
        r.font.size = Pt(9)
        r.font.color.rgb = color
    legend.add_run('\n')

    # 维度五：加梗技术
    r = legend.add_run('  ▸ 加梗技术：')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    punch_items = [
        ("笑点", RGBColor(0xE5, 0x3E, 0x3E)),
        ("呈现", RGBColor(0x66, 0xBB, 0x6A)),
        ("比喻", RGBColor(0x7E, 0x57, 0xC2)),
        ("类比", RGBColor(0x7E, 0x57, 0xC2)),
        ("夸张", RGBColor(0xEC, 0x40, 0x7A)),
        ("假设", RGBColor(0x42, 0xA5, 0xF5)),
        ("混合", RGBColor(0xAB, 0x47, 0xBC)),
        ("call-back", RGBColor(0xF5, 0xA6, 0x23)),
        ("反转", RGBColor(0xAB, 0x47, 0xBC)),
        ("自嘲", RGBColor(0x66, 0xBB, 0x6A)),
        ("金句", RGBColor(0xFF, 0xB3, 0x00)),
        ("递进", RGBColor(0x26, 0xC6, 0xDA)),
    ]
    for text, color in punch_items:
        r = legend.add_run(f' [{text}]')
        r.font.size = Pt(9)
        r.font.color.rgb = color
    legend.add_run('\n')

    # 评估维度说明
    r = legend.add_run('  ▸ AI 评估维度（不标注，嵌入注释）：')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r = legend.add_run(' 核心四要素（好笑/真实/原创/对话）· 素材五特征（真实/个人化/宜人度/熟悉度/普适性）· 精简原则')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 分割线
    doc.add_paragraph('─' * 50)

    # ---- 写入标注段落 ----
    annotated_texts = set()
    for text, tag, note in segments:
        annotated_texts.add(text)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)

        # 技巧标签
        tag_color = get_tag_color(tag)
        tag_run = p.add_run(f'[{tag}]  ')
        tag_run.font.size = Pt(8)
        tag_run.font.bold = True
        tag_run.font.color.rgb = tag_color

        # 正文
        text_run = p.add_run(text)
        text_run.font.size = Pt(11)

        # 注释
        if note:
            note_p = doc.add_paragraph()
            note_p.paragraph_format.left_indent = Cm(1.5)
            note_p.paragraph_format.space_before = Pt(0)
            note_p.paragraph_format.space_after = Pt(6)
            note_r = note_p.add_run(f'💡 {note}')
            note_r.font.size = Pt(9)
            note_r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            note_r.font.italic = True

    # 分割线
    doc.add_paragraph('')
    doc.add_paragraph('─' * 50)

    # ---- 技巧统计（按五维度分类） ----
    doc.add_heading('技巧统计', level=2)

    tag_counts = Counter()
    for _, tag, _ in segments:
        for t in tag.replace('/', '+').replace('·', '+').split('+'):
            t = t.strip()
            if t:
                tag_counts[t] += 1

    # 维度三：结构统计
    struct_tags = ['引子', '论点', '分论点', '荒诞点', '例证', '铺垫', '收尾']
    struct_total = sum(tag_counts.get(t, 0) for t in struct_tags)

    # 维度四：观点统计
    view_tags = ['观点', '负面情绪', '社会观察']
    view_total = sum(tag_counts.get(t, 0) for t in view_tags)

    # 维度五：加梗统计
    punch_tags = ['笑点', 'punchline', '呈现', '比喻', '类比', '夸张', '假设', '混合',
                  'call-back', '回旋镖', '反转', '二次反转', '自嘲', '金句', '递进', '排比',
                  '反讽', 'meta笑点', '解构', '极简收束', '误导', '悬念']
    punch_total = sum(tag_counts.get(t, 0) for t in punch_tags)
    punch_specific = sum(tag_counts.get(t, 0) for t in ['呈现', '比喻', '类比', '夸张', '假设', '混合'])
    cb_total = tag_counts.get('call-back', 0) + tag_counts.get('回旋镖', 0)

    sp = doc.add_paragraph()
    sp.add_run(f'总标注段数：{len(segments)} 段\n\n').font.size = Pt(10)

    r = sp.add_run(f'▸ 结构逻辑：{struct_total} 处')
    r.font.size = Pt(10)
    r.font.bold = True
    sp.add_run(f'  （铺垫 {tag_counts.get("铺垫", 0)} | 例证 {tag_counts.get("例证", 0)} | 荒诞点 {tag_counts.get("荒诞点", 0)}）\n').font.size = Pt(10)

    r = sp.add_run(f'▸ 观点提炼：{view_total} 处')
    r.font.size = Pt(10)
    r.font.bold = True
    sp.add_run(f'  （观点 {tag_counts.get("观点", 0)} | 社会观察 {tag_counts.get("社会观察", 0)}）\n').font.size = Pt(10)

    r = sp.add_run(f'▸ 加梗技术：{punch_total} 处')
    r.font.size = Pt(10)
    r.font.bold = True
    sp.add_run(f'  （呈现 {tag_counts.get("呈现", 0)} | 比喻 {tag_counts.get("比喻", 0)} | 类比 {tag_counts.get("类比", 0)} | 夸张 {tag_counts.get("夸张", 0)} | 假设 {tag_counts.get("假设", 0)} | 混合 {tag_counts.get("混合", 0)}）\n').font.size = Pt(10)
    sp.add_run(f'  └ 反转 {tag_counts.get("反转", 0)} | 自嘲 {tag_counts.get("自嘲", 0)} | call-back {cb_total} | 金句 {tag_counts.get("金句", 0)}\n').font.size = Pt(10)

    # ---- 表演者风格总结（五维度评估） ----
    if performer:
        doc.add_paragraph('')
        doc.add_heading(f'{performer} 风格总结（五维度评估）', level=2)

        # 维度一：核心四要素评估
        doc.add_heading('维度一：核心四要素', level=3)
        four_elem = doc.add_paragraph()
        four_elem.add_run('好笑/真实/原创/对话——由 AI 在标注注释中逐段评估，此处为综合判定。\n').font.size = Pt(10)

        # 维度二：素材选择评估
        doc.add_heading('维度二：素材选择', level=3)
        material = doc.add_paragraph()
        material.add_run('真实性/个人化/宜人度/熟悉度/普适性/荒谬点——由 AI 在标注注释中逐段评估。\n').font.size = Pt(10)

        # 维度三：结构逻辑评估
        doc.add_heading('维度三：结构逻辑', level=3)
        struct_notes = doc.add_paragraph()
        r = struct_notes.add_run(f'结构标签共 {struct_total} 处。')
        r.font.size = Pt(10)
        if tag_counts.get('铺垫', 0) > tag_counts.get('例证', 0):
            struct_notes.add_run(f'铺垫（{tag_counts.get("铺垫", 0)}）多于例证（{tag_counts.get("例证", 0)}），结构偏向"观点先行，例子辅助"的议论文模式。').font.size = Pt(10)
        else:
            struct_notes.add_run(f'例证（{tag_counts.get("例证", 0)}）多于铺垫（{tag_counts.get("铺垫", 0)}），用大量具体例子论证观点。').font.size = Pt(10)
        if tag_counts.get('荒诞点', 0) > 0:
            struct_notes.add_run(f' 荒诞点共 {tag_counts.get("荒诞点", 0)} 处，素材的荒谬性挖掘充分。').font.size = Pt(10)

        # 维度四：观点提炼评估
        doc.add_heading('维度四：观点提炼', level=3)
        view_notes = doc.add_paragraph()
        r = view_notes.add_run(f'观点标签共 {view_total} 处。')
        r.font.size = Pt(10)
        if tag_counts.get('社会观察', 0) > 0:
            view_notes.add_run(f' 社会观察 {tag_counts.get("社会观察", 0)} 处，善于将社会议题转化为喜剧素材。').font.size = Pt(10)
        view_notes.add_run(f'\n观点公式检验：观点 = 事实 + 负面情绪（困难/奇怪/害怕/愚蠢）——逐段注释中已标注。').font.size = Pt(10)

        # 维度五：加梗技术评估
        doc.add_heading('维度五：加梗技术', level=3)
        punch_notes = doc.add_paragraph()
        r = punch_notes.add_run(f'加梗标签共 {punch_total} 处。')
        r.font.size = Pt(10)

        # 找出最常用的加梗技术
        tech_counts = [(t, tag_counts[t]) for t in punch_tags if tag_counts.get(t, 0) > 0]
        tech_counts.sort(key=lambda x: x[1], reverse=True)
        if tech_counts:
            top3 = tech_counts[:3]
            top_str = ' > '.join([f'{t}({c})' for t, c in top3])
            punch_notes.add_run(f'\n最常用技术：{top_str}').font.size = Pt(10)

        if cb_total > 0:
            punch_notes.add_run(f'\ncall-back {cb_total} 次，首尾呼应意识强。').font.size = Pt(10)
        if tag_counts.get('自嘲', 0) > 0:
            punch_notes.add_run(f'\n自嘲 {tag_counts.get("自嘲", 0)} 次，用"把自己放进去"建立亲和力。').font.size = Pt(10)
        if tag_counts.get('金句', 0) > 0:
            punch_notes.add_run(f'\n金句 {tag_counts.get("金句", 0)} 处，点睛能力突出。').font.size = Pt(10)

    # ---- 纯文本版 ----
    doc.add_page_break()
    doc.add_heading('纯文本版', level=2)
    doc.add_paragraph('')
    full_text = '\n\n'.join(lines)
    doc.add_paragraph(full_text)

    # 保存
    doc.save(output_path)
    print(f'DONE: {output_path}')
    return output_path


# ============================================================
# 命令行入口
# ============================================================

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='喜剧技巧标注生成器')
    parser.add_argument('input', help='输入文件（JSON 或 文本文件）')
    parser.add_argument('output', nargs='?', default=None, help='输出 Word 路径')
    parser.add_argument('--performer', default='', help='表演者名字')
    parser.add_argument('--title', default='喜剧技巧标注版', help='文档标题')
    args = parser.parse_args()

    # 读取输入
    with open(args.input, 'r', encoding='utf-8') as f:
        raw = f.read().strip()

    # 尝试解析为 JSON（Whisper segments）
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            text_lines = data
        else:
            text_lines = [raw]
    except json.JSONDecodeError:
        text_lines = [l for l in raw.split('\n') if l.strip()]

    generate_annotated_word(
        text_lines=text_lines,
        title=args.title,
        output_path=args.output,
        performer=args.performer,
    )
