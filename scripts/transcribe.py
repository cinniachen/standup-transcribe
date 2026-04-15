#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
脱口秀视频转录工具
支持：mp4, ts, mkv, avi, mov, m4a, mp3, wav 等格式
功能：Whisper 转录 → 自动校对人名/关键词 → 生成带时间戳字幕 + 纯文本 Word
"""

import os
import sys
import shutil
import tempfile
import subprocess
import json

# ---------- 配置 ----------
FFMPEG_PATH = shutil.which("ffmpeg")
WHISPER_MODEL = "base"  # tiny/base/small/turbo/large-v3
LANGUAGE = "zh"

# ---------- 校对相关 ----------
# names.json 路径：与脚本同目录
_NAMES_JSON = os.path.join(os.path.dirname(os.path.abspath(__file__)), "names.json")

def load_corrections(extra_corrections=None):
    """
    延迟加载校对字典：从 names.json 读取别名映射 + 内置关键词。
    只在校对阶段调用，不浪费启动时的 token/内存。
    :param extra_corrections: 用户额外传入的 {错误: 正确} 字典
    :return: 合并后的校对字典
    """
    corrections = {}
    # 1. 从 names.json 加载别名映射
    if os.path.isfile(_NAMES_JSON):
        with open(_NAMES_JSON, "r", encoding="utf-8") as f:
            names_data = json.load(f)
        corrections.update(names_data.get("易混淆别名映射", {}))

    # 2. 内置关键词校对（Whisper 常见错误）
    corrections.update({
        "脱口修": "脱口秀",
        "搭谝": "搭讪",
        "仙人长": "仙人掌",
        "金钥带": "金腰带",
        "四角笼": "八角笼",
    })

    # 3. 节目名标准化
    corrections.update({
        "脱口秀和ta的朋友们": "脱口秀和TA的朋友们",
        "主咖和ta的朋友们": "主咖和TA的朋友们",
    })

    # 4. 用户自定义覆盖
    if extra_corrections:
        corrections.update(extra_corrections)

    return corrections


def find_ffmpeg():
    """查找 ffmpeg"""
    candidates = [
        os.path.expanduser("~/FFmpeg/ffmpeg"),
        "/usr/local/bin/ffmpeg",
        "/opt/homebrew/bin/ffmpeg",
    ]
    system_ffmpeg = shutil.which("ffmpeg")
    if system_ffmpeg:
        candidates.insert(0, system_ffmpeg)
    for p in candidates:
        if p and os.path.isfile(p):
            return p
    return None


def extract_audio(video_path, audio_path, ffmpeg):
    """从视频中提取 16kHz 单声道 wav 音频"""
    cmd = [
        ffmpeg, "-i", video_path,
        "-vn", "-acodec", "pcm_s16le",
        "-ar", "16000", "-ac", "1",
        "-y", audio_path
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=300)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg 提取音频失败: {result.stderr.decode()[:200]}")
    return audio_path


def transcribe(audio_path, model_name=WHISPER_MODEL, language=LANGUAGE):
    """Whisper 转录"""
    # 确保 whisper 能找到 ffmpeg
    ffmpeg_dir = os.path.dirname(find_ffmpeg()) if find_ffmpeg() else ""
    if ffmpeg_dir and ffmpeg_dir not in os.environ.get("PATH", ""):
        os.environ["PATH"] = ffmpeg_dir + ":" + os.environ.get("PATH", "")
    import whisper
    model = whisper.load_model(model_name)
    result = model.transcribe(audio_path, language=language, verbose=False)
    return result


def correct_text(text, corrections):
    """自动校对文本"""
    for wrong, right in corrections.items():
        if wrong != right:
            text = text.replace(wrong, right)
    return text


def generate_worddoc(segments, title, output_path, corrections=None):
    """生成 Word 文档（带时间戳 + 纯文本版）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "PingFang SC"
    doc.styles["Normal"].font.size = Pt(11)

    # 标题
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Whisper AI 转录 · 自动校对")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")

    # 带时间戳版本
    for seg in segments:
        s, e = seg["start"], seg["end"]
        ts = f"[{int(s//60):02d}:{int(s%60):02d} - {int(e//60):02d}:{int(e%60):02d}]"
        text = correct_text(seg["text"].strip(), corrections)
        p = doc.add_paragraph()
        tr = p.add_run(ts + "  ")
        tr.font.size = Pt(9)
        tr.font.color.rgb = RGBColor(100, 100, 100)
        pr = p.add_run(text)
        pr.font.size = Pt(11)

    # 纯文本版
    doc.add_page_break()
    doc.add_heading("纯文本版", level=2)
    doc.add_paragraph("")
    full = "\n\n".join(correct_text(seg["text"].strip(), corrections) for seg in segments)
    doc.add_paragraph(full)

    doc.save(output_path)
    return output_path


def process_single(video_path, output_dir=None, title=None, corrections=None, model=None):
    """
    处理单个视频文件
    :param video_path: 视频文件路径
    :param output_dir: 输出目录（默认与视频同目录）
    :param title: 文档标题（默认用文件名）
    :param corrections: 自定义校对字典 {错误: 正确}（会与 names.json 合并）
    :param model: whisper 模型名
    :return: {"output": 输出文件路径, "segments": 转录段列表}
    """
    video_path = os.path.expanduser(video_path)
    if not os.path.isfile(video_path):
        raise FileNotFoundError(f"文件不存在: {video_path}")

    output_dir = output_dir or os.path.dirname(video_path) or os.path.expanduser("~/Downloads")
    os.makedirs(output_dir, exist_ok=True)

    basename = os.path.splitext(os.path.basename(video_path))[0]
    title = title or basename
    model = model or WHISPER_MODEL

    ffmpeg = find_ffmpeg()
    if not ffmpeg:
        raise RuntimeError("未找到 ffmpeg，请先安装：brew install ffmpeg")

    tmpdir = tempfile.mkdtemp(prefix="standup_transcribe_")
    try:
        # 1. 提取音频
        audio_path = os.path.join(tmpdir, "audio.wav")
        print(f"[1/3] 提取音频: {video_path}")
        extract_audio(video_path, audio_path, ffmpeg)
        print(f"      ✅ 完成")

        # 2. Whisper 转录
        print(f"[2/3] Whisper 转录 (model={model})...")
        result = transcribe(audio_path, model_name=model)
        segments = result["segments"]
        duration = segments[-1]["end"] if segments else 0
        print(f"      ✅ 完成 ({len(segments)} 段, {duration:.0f}秒)")

        # 3. 加载校对字典并生成 Word（仅在此阶段读取 names.json）
        cor = load_corrections(corrections)
        output_path = os.path.join(output_dir, f"{basename}_字幕.docx")
        print(f"[3/3] 生成 Word（校对字典 {len(cor)} 条）...")
        generate_worddoc(segments, title, output_path, cor)
        print(f"      ✅ 完成: {output_path}")

        return {"output": output_path, "segments": segments}

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def process_batch(video_paths, output_dir=None, corrections=None, model=None):
    """
    批量处理多个视频文件
    :param video_paths: 视频文件路径列表
    :param output_dir: 输出目录
    :param corrections: 自定义校对字典
    :param model: whisper 模型名
    :return: 成功的文件列表
    """
    results = []
    total = len(video_paths)

    print(f"\n{'='*60}")
    print(f"  批量转录: 共 {total} 个文件")
    print(f"{'='*60}\n")

    for i, vp in enumerate(video_paths, 1):
        print(f"\n{'─'*40}")
        print(f"  [{i}/{total}] {os.path.basename(vp)}")
        print(f"{'─'*40}")
        try:
            out = process_single(vp, output_dir=output_dir, corrections=corrections, model=model)
            results.append({"file": vp, "output": out["output"], "segments": out.get("segments", []), "status": "ok"})
        except Exception as e:
            print(f"      ❌ 失败: {e}")
            results.append({"file": vp, "output": None, "status": f"error: {e}"})

    # 汇总
    print(f"\n{'='*60}")
    print(f"  批量转录完成")
    ok = sum(1 for r in results if r["status"] == "ok")
    print(f"  成功: {ok}/{total}")
    for r in results:
        icon = "✅" if r["status"] == "ok" else "❌"
        print(f"    {icon} {os.path.basename(r['file'])} → {os.path.basename(r['output']) if r['output'] else r['status']}")
    print(f"{'='*60}")

    return results


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法:")
        print(f"  python3 {sys.argv[0]} 视频文件1 [视频文件2 ...]")
        print(f"  python3 {sys.argv[0]} ~/Downloads/*.mp4")
        print(f"  python3 {sys.argv[0]} --output ~/Documents 视频文件")
        print("")
        print("支持格式: mp4, ts, mkv, avi, mov, m4a, mp3, wav")
        sys.exit(0)

    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    output_dir = None
    for a in sys.argv[1:]:
        if a in ("--output", "-o") or a.startswith("--output="):
            idx = sys.argv.index(a)
            if "=" in a:
                output_dir = a.split("=", 1)[1]
            elif idx + 1 < len(sys.argv):
                output_dir = sys.argv[idx + 1]
                args.remove(output_dir)

    process_batch(args, output_dir=output_dir)
